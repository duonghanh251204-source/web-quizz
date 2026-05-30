"""
document_processor.py
Trích xuất text từ PDF, DOCX, TXT và chia nhỏ thành các chunk có metadata.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

from loguru import logger


@dataclass
class DocumentChunk:
    """Một đoạn văn bản với metadata đầy đủ."""
    content: str
    page: Optional[int] = None
    chapter: Optional[str] = None
    source_file: str = ""
    chunk_index: int = 0
    total_chunks: int = 0


@dataclass
class ProcessedDocument:
    chunks: List[DocumentChunk] = field(default_factory=list)
    title: str = ""
    total_pages: int = 0
    source_file: str = ""
    extraction_method: str = ""


# ---------------------------------------------------------------------------
# Text extraction helpers
# ---------------------------------------------------------------------------

def _extract_pdf(file_bytes: bytes, filename: str) -> ProcessedDocument:
    """Trích xuất PDF, fallback sang OCR nếu cần."""
    import pymupdf4llm  # type: ignore
    import fitz          # PyMuPDF

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    total_pages = doc.page_count

    # Thử pymupdf4llm để giữ cấu trúc Markdown (heading, bảng)
    try:
        md_text = pymupdf4llm.to_markdown(doc)
        raw_text = md_text
        method = "pymupdf4llm"
        logger.info(f"[PDF] Dùng pymupdf4llm — {len(raw_text)} ký tự")
    except Exception as e:
        logger.warning(f"[PDF] pymupdf4llm lỗi ({e}), fallback sang pdfplumber")
        raw_text, method = _extract_pdf_pdfplumber(file_bytes)

    # Nếu text quá ít → PDF dạng scan, dùng OCR
    if len(raw_text.strip()) < 100:
        logger.warning("[PDF] Ít text, thử OCR")
        raw_text, method = _extract_pdf_ocr(file_bytes)

    return ProcessedDocument(
        title=Path(filename).stem,
        total_pages=total_pages,
        source_file=filename,
        extraction_method=method,
        chunks=[]  # chunks được điền sau khi chunk
    ), raw_text


def _extract_pdf_pdfplumber(file_bytes: bytes):
    import pdfplumber  # type: ignore
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            t = page.extract_text() or ""
            text_parts.append(f"[Trang {i+1}]\n{t}")
    return "\n\n".join(text_parts), "pdfplumber"


def _extract_pdf_ocr(file_bytes: bytes):
    try:
        from pdf2image import convert_from_bytes  # type: ignore
        import pytesseract                          # type: ignore
        images = convert_from_bytes(file_bytes)
        text_parts = []
        for i, img in enumerate(images):
            t = pytesseract.image_to_string(img, lang="vie+eng")
            text_parts.append(f"[Trang {i+1}]\n{t}")
        return "\n\n".join(text_parts), "OCR (pytesseract)"
    except Exception as e:
        logger.error(f"[OCR] Thất bại: {e}")
        return "", "OCR failed"


def _extract_docx(file_bytes: bytes, filename: str):
    from docx import Document  # type: ignore
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    raw_text = "\n\n".join(parts)
    title = doc.core_properties.title or Path(filename).stem
    return ProcessedDocument(
        title=title,
        total_pages=0,
        source_file=filename,
        extraction_method="python-docx",
        chunks=[]
    ), raw_text


def _extract_txt(file_bytes: bytes, filename: str):
    raw_text = file_bytes.decode("utf-8", errors="ignore")
    return ProcessedDocument(
        title=Path(filename).stem,
        total_pages=0,
        source_file=filename,
        extraction_method="plain-text",
        chunks=[]
    ), raw_text


# ---------------------------------------------------------------------------
# Text cleaning
# ---------------------------------------------------------------------------

def _clean_text(text: str) -> str:
    text = re.sub(r'\n{3,}', '\n\n', text)           # Nhiều dòng trắng → 2
    text = re.sub(r'[ \t]{2,}', ' ', text)            # Nhiều khoảng trắng → 1
    text = re.sub(r'\x00', '', text)                   # Null bytes
    return text.strip()


# ---------------------------------------------------------------------------
# Chunking (RecursiveCharacterTextSplitter style)
# ---------------------------------------------------------------------------

def _chunk_text(
    text: str,
    chunk_size: int = 2000,
    chunk_overlap: int = 200,
) -> List[str]:
    """Chia text thành các chunk có overlap."""
    separators = ["\n\n", "\n", ". ", " ", ""]
    chunks = _recursive_split(text, separators, chunk_size, chunk_overlap)
    return [c.strip() for c in chunks if c.strip()]


def _recursive_split(text: str, separators: List[str], chunk_size: int, overlap: int) -> List[str]:
    if not separators:
        # Chia theo ký tự
        return [text[i:i+chunk_size] for i in range(0, len(text), chunk_size - overlap)]

    sep = separators[0]
    splits = text.split(sep) if sep else list(text)
    chunks: List[str] = []
    current = ""

    for split in splits:
        candidate = (current + sep + split) if current else split
        if len(candidate) <= chunk_size:
            current = candidate
        else:
            if current:
                chunks.append(current)
            if len(split) > chunk_size:
                # Đệ quy với separator tiếp theo
                sub = _recursive_split(split, separators[1:], chunk_size, overlap)
                chunks.extend(sub)
                current = sub[-1][-overlap:] if sub else ""
            else:
                current = split

    if current:
        chunks.append(current)
    return chunks


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def process_document(
    file_bytes: bytes,
    filename: str,
    chunk_size: int = 2000,
    chunk_overlap: int = 200,
) -> ProcessedDocument:
    """
    Điểm vào chính: nhận bytes của file → trả về ProcessedDocument với đầy đủ chunks.
    """
    ext = Path(filename).suffix.lower()

    logger.info(f"Xử lý file: {filename} (ext={ext})")

    if ext == ".pdf":
        doc, raw_text = _extract_pdf(file_bytes, filename)
    elif ext in (".docx", ".doc"):
        doc, raw_text = _extract_docx(file_bytes, filename)
    elif ext == ".txt":
        doc, raw_text = _extract_txt(file_bytes, filename)
    else:
        raise ValueError(f"Định dạng không hỗ trợ: {ext}")

    raw_text = _clean_text(raw_text)
    text_chunks = _chunk_text(raw_text, chunk_size, chunk_overlap)

    doc.chunks = [
        DocumentChunk(
            content=chunk,
            source_file=filename,
            chunk_index=i,
            total_chunks=len(text_chunks),
        )
        for i, chunk in enumerate(text_chunks)
    ]

    logger.info(f"Hoàn thành: {len(doc.chunks)} chunks từ '{filename}'")
    return doc
